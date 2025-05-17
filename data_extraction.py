import os
import logging
import pytesseract
from PIL import Image
import cv2
import PyPDF2
import docx
import pandas as pd
from datetime import datetime
import tempfile
import html2text
from email import policy
from email.parser import BytesParser
from docx2pdf import convert as docx_to_pdf
from google.cloud import vision
import base64
import requests
import json

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Adjust path as needed


# --- Logging setup
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# --- Utilities
def get_file_extension(file_path):
    return os.path.splitext(file_path)[-1].replace('.', '').lower()

def perform_ocr(file_path):
    try:
        # Read the image file
        with open(file_path, "rb") as image_file:
            content = image_file.read()

        # Create the payload for the Vision API request
        request_body = {
            "requests": [
                {
                    "image": {
                        "content": base64.b64encode(content).decode("utf-8")
                    },
                    "features": [
                        {"type": "TEXT_DETECTION"}
                    ]
                }
            ]
        }

        # Use your API key
        api_key = "your_api_key"
        
        # Make the POST request
        response = requests.post(
            f"https://vision.googleapis.com/v1/images:annotate?key={api_key}",
            json=request_body
        )

        # Process the response
        if response.status_code == 200:
            result = response.json()
            text_annotations = result.get("responses", [])[0].get("textAnnotations", [])
            
            if text_annotations:
                full_text = text_annotations[0].get("description", "")
                return full_text
            else:
                return "No text detected."
        else:
            logger.error(f"OCR API Error: {response.status_code} {response.text}")
            return f"OCR API Error: {response.status_code} {response.text}"
    
    except Exception as e:
        logger.error(f"OCR failed: {str(e)}")
        return f"OCR failed: {str(e)}"

def extract_text_from_word(file_path):
    logger.debug(f"Extracting text from Word document: {file_path}")
    try:
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join([cell.text for cell in row.cells])
        return text
    except Exception as e:
        logger.error(f"Word extraction error: {str(e)}")
        return convert_docx_to_pdf_and_extract(file_path)

def extract_text_from_pdf(file_path):
    logger.debug(f"Extracting text from PDF: {file_path}")
    try:
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            if reader.is_encrypted:
                raise ValueError("PDF is encrypted.")
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if not page_text.strip():
                    logger.debug(f"No text on page {i}, fallback OCR.")
                    page_text = perform_ocr(file_path)
                text += page_text + "\n\n"
        return text
    except Exception as e:
        logger.error(f"PDF processing error: {str(e)}")
        return perform_ocr(file_path)

def extract_text_from_image(file_path):
    logger.debug(f"Extracting text from image: {file_path}")
    return perform_ocr(file_path)

def convert_docx_to_pdf_and_extract(file_path):
    with tempfile.TemporaryDirectory() as temp_dir:
        pdf_path = os.path.join(temp_dir, "converted.pdf")
        try:
            docx_to_pdf(file_path, pdf_path)
            return extract_text_from_pdf(pdf_path)
        except Exception as e:
            raise RuntimeError(f"Failed to convert .docx to PDF: {e}")

def load_email_text(file_path):
    with open(file_path, 'rb') as f:
        msg = BytesParser(policy=policy.default).parse(f)

    body = msg.get_body(preferencelist=('plain', 'html'))
    if body:
        try:
            content = body.get_content()
            return html2text.html2text(content) if body.get_content_type() == 'text/html' else content
        except Exception as e:
            logger.warning(f"Email body read error: {e}")

    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() in ['text/plain', 'text/html'] and part.get_content_disposition() != 'attachment':
                try:
                    content = part.get_content()
                    return html2text.html2text(content) if part.get_content_type() == 'text/html' else content
                except Exception as e:
                    logger.warning(f"Email part read error: {e}")
    else:
        try:
            content = msg.get_content()
            return html2text.html2text(content) if msg.get_content_type() == 'text/html' else content
        except Exception as e:
            logger.warning(f"Email read error: {e}")

    return "⚠️ No readable content found in the email."

def extract_text_from_excel(file_path):
    try:
        df = pd.read_excel(file_path, sheet_name=None)
        combined_text = ""
        for sheet_name, sheet in df.items():
            combined_text += f"\n--- Sheet: {sheet_name} ---\n"
            combined_text += sheet.to_string(index=False)
        return combined_text
    except Exception as e:
        logger.error(f"Excel processing error: {str(e)}")
        return f"Excel error: {e}"

# --- Dispatcher
def process_document(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_ext = get_file_extension(file_path)
    logger.info(f"Detected file type: .{file_ext}")

    if file_ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext in ['docx', 'doc']:
        return extract_text_from_word(file_path)
    elif file_ext in ['png', 'jpg', 'jpeg', 'tiff']:
        return extract_text_from_image(file_path)
    elif file_ext in ['eml']:
        return load_email_text(file_path)
    elif file_ext in ['xlsx', 'xls']:
        return extract_text_from_excel(file_path)
    else:
        raise ValueError(f"Unsupported file type: .{file_ext}")

# --- Batch Processor
if __name__ == "__main__":
    input_folder = "test data/Termsheets"
    output_folder = "extracted_files"
    os.makedirs(output_folder, exist_ok=True)

    for filename in os.listdir(input_folder):
        if not filename.lower().endswith((".pdf", ".docx", ".doc", ".eml", ".xlsx", ".xls", ".jpg", ".jpeg", ".png", ".tiff")):
            continue

        file_path = os.path.join(input_folder, filename)
        print(f"\n📄 Processing: {file_path}")

        try:
            content = process_document(file_path)
            base_name = os.path.splitext(filename)[0]
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            output_filename = f"{base_name}_extracted_{timestamp}.txt"
            output_path = os.path.join(output_folder, output_filename)

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)

            print(f"✅ Saved to: {output_path}")
        except Exception as e:
            print(f"❌ Error processing {filename}: {e}")
